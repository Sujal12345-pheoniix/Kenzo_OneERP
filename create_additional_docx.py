import os
import shutil
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def build_diagrams_docx():
    doc = docx.Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)
        
    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_t.add_run("KENZO ONEERP (KORE) - ARCHITECTURE & DIAGRAMS HANDBOOK")
    r_t.bold = True
    r_t.font.name = 'Segoe UI'
    r_t.font.size = Pt(22)
    r_t.font.color.rgb = RGBColor(67, 56, 202)
    
    p_s = doc.add_paragraph()
    p_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_s.paragraph_format.space_after = Pt(18)
    r_sub = p_s.add_run("Complete Collection of Visual System Architecture, DFD Level 1, ERD Schema, RBAC Pyramid, and Theme Engine Diagrams")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = RGBColor(100, 116, 139)

    diagrams = [
        ("1. Full-Stack System Architecture Diagram", "docx_assets/arch_diagram.png", 
         "Explains the end-to-end data flow from Client Browser -> Next.js App Router -> Session Auth -> Prisma ORM -> Neon PostgreSQL DB & Cloudinary CDN."),
        ("2. Data Flow Diagram (DFD Level 1)", "docx_assets/dfd_level1.png", 
         "Illustrates the real-time submission of leave applications, database write locks, and real-time polling/notifications on Admin/CEO/HR dashboards."),
        ("3. Leave State Lifecycle & Isolation Workflow", "docx_assets/leave_workflow.png", 
         "Details the 4-step leave lifecycle: Employee submission, tenant-scoped PENDING status, role-based visibility filter, and APPROVED/REJECTED status update."),
        ("4. Database Entity Relationship Diagram (ERD)", "docx_assets/erd_diagram.png", 
         "Visualizes Prisma ORM relational database models: Tenant, User, Employee, Attendance, Leave, Project, Task, Invoice, Expense."),
        ("5. Role-Based Access Control (RBAC) Hierarchy", "docx_assets/rbac_hierarchy.png", 
         "Pyramid breakdown of user roles from SUPER_ADMIN & COMPANY_ADMIN down to CEO/CTO, HR/FINANCE, and DEVELOPER/EMPLOYEE levels."),
        ("6. Theme Engine & Text Contrast Architecture", "docx_assets/theme_engine.png", 
         "Shows how CSS contrast variables manage Light Mode (#0f172a text) vs Dark Mode (#ffffff text) and fixed modal styling (.profile-modal-dark).")
    ]

    for title, img_path, desc in diagrams:
        p_h = doc.add_paragraph()
        p_h.paragraph_format.space_before = Pt(16)
        p_h.paragraph_format.space_after = Pt(4)
        r_h = p_h.add_run(title)
        r_h.bold = True
        r_h.font.name = 'Segoe UI'
        r_h.font.size = Pt(14)
        r_h.font.color.rgb = RGBColor(2, 132, 199)

        p_d = doc.add_paragraph()
        p_d.paragraph_format.space_after = Pt(8)
        r_d = p_d.add_run(desc)
        r_d.font.name = 'Calibri'
        r_d.font.size = Pt(11)
        r_d.font.color.rgb = RGBColor(51, 65, 85)

        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(6.2))
            doc.add_paragraph().paragraph_format.space_after = Pt(12)

    doc.save("Kenzo_OneERP_Diagrams_Handbook.docx")
    print("Saved Kenzo_OneERP_Diagrams_Handbook.docx")

def build_presentation_docx():
    doc = docx.Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_t.add_run("KENZO ONEERP - SYSTEM PRESENTATION DECK")
    r_t.bold = True
    r_t.font.name = 'Segoe UI'
    r_t.font.size = Pt(22)
    r_t.font.color.rgb = RGBColor(67, 56, 202)

    p_s = doc.add_paragraph()
    p_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_s.paragraph_format.space_after = Pt(18)
    r_sub = p_s.add_run("Executive Presentation Slides Transcripts & System Upgrade Highlights")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = RGBColor(100, 116, 139)

    slides_content = [
        ("Slide 1: Executive Overview", [
            ("Title: ", "Kenzo OneERP - Advanced Employee Portals & Smart Attendance Systems"),
            ("Tagline: ", "Next-Generation Enterprise Resource Planning Suite"),
            ("Summary: ", "An optimized enterprise resource system upgrading operational transparency with interactive admin profile portals, role-based access control, and cryptographically verified date check-in calendars.")
        ]),
        ("Slide 2: Current Operational Challenges", [
            ("Siloed Employee Data: ", "Management had to jump across departments and payrolls to inspect team member output."),
            ("Attendance Back-dating & Leakage: ", "Traditional timesheets allowed users to check in retrospectively for past dates or prepopulate future days."),
            ("Admin Overheads: ", "Administrators spent hours aggregating metrics manually.")
        ]),
        ("Slide 3: Strategic Rationale & Solution", [
            ("Operational Auditability: ", "Guarantees a clean, unmanipulated paper trail of attendance records."),
            ("Resource Synchronization: ", "Links employee tasks, leaves, and payroll directly to their profiles."),
            ("Frictionless Compliance: ", "Automates policies like date boundaries, check-in thresholds, and role exemptions.")
        ]),
        ("Slide 4: Core Module Upgrades", [
            ("Admin Profile Hub: ", "Unified profile dashboard with project container tracking, task boards, and historical attendance grids."),
            ("Strict Check-In Calendar: ", "Employees can only check in for the current calendar date; past and future days are locked."),
            ("Executive Exemptions: ", "Admins and CEOs are exempted from daily check-in prompt clutter.")
        ]),
        ("Slide 5: Technology Stack", [
            ("Frontend: ", "Next.js App Router, React 19, TypeScript, Tailwind CSS."),
            ("Data Layer: ", "Prisma ORM v5, Neon PostgreSQL Cloud Database."),
            ("Security & Media: ", "JWT tokens, BcryptJS password hashing, Cloudinary CDN signed uploads.")
        ]),
        ("Slide 6: System Flow & Verification Model", [
            ("Step 1: ", "Request Session verification via JWT & role checking."),
            ("Step 2: ", "Role verification check (Is Admin/CEO?)."),
            ("Step 3: ", "Date boundary check (Compare check-in date with Today's UTC date)."),
            ("Step 4: ", "Database write lock enforcing @@unique([employeeId, date]).")
        ]),
        ("Slide 7: Key Advantages", [
            ("Anti-Tamper Design: ", "Removes backdated and pre-filled timesheets completely."),
            ("Executive Clarity: ", "Eliminates visual noise for Admins and CEOs."),
            ("Real-Time Analytics: ", "Syncs attendance logs with ongoing project task Kanban charts.")
        ])
    ]

    for title, items in slides_content:
        p_h = doc.add_paragraph()
        p_h.paragraph_format.space_before = Pt(14)
        p_h.paragraph_format.space_after = Pt(4)
        r_h = p_h.add_run(title)
        r_h.bold = True
        r_h.font.name = 'Segoe UI'
        r_h.font.size = Pt(14)
        r_h.font.color.rgb = RGBColor(2, 132, 199)

        for b_prefix, text in items:
            p_b = doc.add_paragraph(style='List Bullet')
            p_b.paragraph_format.space_after = Pt(3)
            r_bp = p_b.add_run(b_prefix)
            r_bp.bold = True
            r_bp.font.name = 'Calibri'
            r_bp.font.size = Pt(11)
            r_bp.font.color.rgb = RGBColor(15, 23, 42)

            r_tx = p_b.add_run(text)
            r_tx.font.name = 'Calibri'
            r_tx.font.size = Pt(11)
            r_tx.font.color.rgb = RGBColor(51, 65, 85)

    doc.save("Kenzo_OneERP_Presentation_Deck.docx")
    print("Saved Kenzo_OneERP_Presentation_Deck.docx")

if __name__ == "__main__":
    build_diagrams_docx()
    build_presentation_docx()
