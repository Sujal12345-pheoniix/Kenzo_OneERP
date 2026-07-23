import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = docx.Document()
    
    # Page Setup (Margins)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Helper Functions for Styling
    def set_cell_background(cell, fill_hex):
        tcPr = cell._element.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)
        
    def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
        tcPr = cell._element.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    def add_callout(text, title="KEY DEVELOPER INSIGHT"):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_background(cell, "F0F9FF") # Light sky blue
        set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
        
        # Left border blue
        tcPr = cell._element.get_or_add_tcPr()
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="0284C7"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
        tcPr.append(borders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r_title = p.add_run(f"💡 {title}\n")
        r_title.bold = True
        r_title.font.name = 'Calibri'
        r_title.font.size = Pt(10.5)
        r_title.font.color.rgb = RGBColor(3, 105, 161)
        
        r_text = p.add_run(text)
        r_text.font.name = 'Calibri'
        r_text.font.size = Pt(10)
        r_text.font.color.rgb = RGBColor(15, 23, 42)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def add_code_block(code_text, language="TypeScript"):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_background(cell, "0F172A") # Dark slate
        set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
        
        tcPr = cell._element.get_or_add_tcPr()
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="0" w:color="6366F1"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
        tcPr.append(borders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        
        r_lang = p.add_run(f"// --- Code Snippet: {language} ---\n")
        r_lang.font.name = 'Consolas'
        r_lang.font.size = Pt(9)
        r_lang.font.color.rgb = RGBColor(148, 163, 184)
        
        r_code = p.add_run(code_text)
        r_code.font.name = 'Consolas'
        r_code.font.size = Pt(9.5)
        r_code.font.color.rgb = RGBColor(248, 250, 252)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Segoe UI'
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(67, 56, 202) # Indigo 700
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Segoe UI'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(2, 132, 199) # Sky 600
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Segoe UI'
        run.font.size = Pt(11.5)
        run.font.color.rgb = RGBColor(124, 58, 237) # Violet 600
        return p

    def add_p(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(15, 23, 42)
        return p

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        r_b = p.add_run(bold_prefix)
        r_b.bold = True
        r_b.font.name = 'Calibri'
        r_b.font.size = Pt(10.5)
        r_b.font.color.rgb = RGBColor(15, 23, 42)
        
        r_t = p.add_run(text)
        r_t.font.name = 'Calibri'
        r_t.font.size = Pt(10.5)
        r_t.font.color.rgb = RGBColor(51, 65, 85)
        return p

    # --- TITLE / HEADER BANNER ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(24)
    p_title.paragraph_format.space_after = Pt(4)
    r_t = p_title.add_run("KENZO ONEERP (KORE)")
    r_t.bold = True
    r_t.font.name = 'Segoe UI'
    r_t.font.size = Pt(26)
    r_t.font.color.rgb = RGBColor(67, 56, 202)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(18)
    r_s = p_sub.add_run("The Complete Solo Full-Stack Developer Architecture & Implementation Masterclass\n(Hinglish Practical Edition)")
    r_s.bold = True
    r_s.font.name = 'Calibri'
    r_s.font.size = Pt(14)
    r_s.font.color.rgb = RGBColor(71, 85, 105)

    # Divider line
    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_div.paragraph_format.space_after = Pt(18)
    p_div.add_run("═════════════════════════════════════════════════════════════════════").font.color.rgb = RGBColor(203, 213, 225)

    # --- CHAPTER 1: OVERVIEW & PRODUCT VISION ---
    add_h1("CHAPTER 1: Project Overview & System Philosophy (ERP Concept)")
    
    add_p("Welcome to the Masterclass Guide for Kenzo OneERP (KORE)! Agar aapne yeh poora ERP project zero se single-handedly build kiya hai, ya aap iski internal working ko completely master karna chahte ho, toh yeh document aapka final blueprint hai. Is guide mein hum visual diagrams, architecture charts, code snippets, aur in-depth Hinglish explanations ke saath har single function, database table, security rule, aur business logic ko dissect karenge.")
    
    add_callout(
        "Kenzo OneERP (KORE) ek Enterprise-grade Multi-Tenant ERP system hai jisme Single Codebase aur Single Database Multi-Tenant Architecture ka upyog kiya gaya hai. Isme Har tenant (Company) ka data logical multi-tenancy rules (tenantId cascading) ke zariye completely isolated aur secure rehta hai.",
        "CORE ARCHITECTURAL PHILOSOPHY"
    )

    add_h2("1.1 Core Business Modules Matrix")
    add_p("Kenzo OneERP core business operations ko single dashboard hub mein unify karta hai. Niche iske 8 primary modules diye gaye hain:")

    # Table of Modules
    t_mod = doc.add_table(rows=1, cols=3)
    t_mod.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t_mod.rows[0].cells
    hdr[0].text = "Module Name"
    hdr[1].text = "Primary Business Purpose"
    hdr[2].text = "Key Operations & Logic"
    
    for i in range(3):
        set_cell_background(hdr[i], "0F172A")
        set_cell_margins(hdr[i], top=100, bottom=100, left=150, right=150)
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        hdr[i].paragraphs[0].runs[0].font.bold = True

    modules_data = [
        ("HRMS & Attendance", "Staff directory & daily attendance tracking", "Check-in/out timestamps, present/late/absent status indexing, monthly attendance history."),
        ("Leave Engine", "Real-time leave application & approval pipeline", "Employee self-application, auto duration calculation, real-time approval/rejection board for Admin/CEO/HR, self-deletion isolation."),
        ("CRM Pipeline", "Sales leads & customer relationship management", "Lead capture, funnel conversion stages (NEW, CONTACTED, QUALIFIED, WON), lead value calculation in ₹."),
        ("Finance & Invoices", "Corporate billing, expenses & revenue ledger", "Invoice generation, status tracking (DRAFT, SENT, PAID, OVERDUE), operational cash inflow calculation, expense approval board."),
        ("Projects & Kanban", "Initiative scope & task execution board", "Project container creation, budget allocation, Kanban drag-and-drop task assignment (TODO, IN_PROGRESS, REVIEW, DONE), priority flags."),
        ("Departments", "Departmental headcount & salary analytics", "Aggregate headcount, annual payroll calculations, average salary indexing, team member avatar cards."),
        ("AI Copilot Hub", "Natural language ERP query assistant", "AI preset queries, automated project roadmap task drafting, financial margin summaries."),
        ("Profile & Identity", "User identity, CDN storage & RBAC profile modal", "Cloudinary CDN base64 upload, signed SHA1 image upload, user-scoped profile isolation (kore_user_profile_${user.id}).")
    ]

    for mod_name, purpose, logic in modules_data:
        row = t_mod.add_row().cells
        row[0].text = mod_name
        row[1].text = purpose
        row[2].text = logic
        for j in range(3):
            set_cell_margins(row[j], top=80, bottom=80, left=120, right=120)
            row[j].paragraphs[0].runs[0].font.size = Pt(9.5)
            row[j].paragraphs[0].runs[0].font.name = 'Calibri'

    # --- CHAPTER 2: TECH STACK & SYSTEM ARCHITECTURE ---
    add_h1("CHAPTER 2: Tech Stack & System Architecture Overview")
    
    add_p("KORE ko modern web application architecture ke best practices par construct kiya gaya hai. Speed, security, maintainability, aur aesthetic design system ko balance karne ke liye following modern tech stack adopt kiya gaya:")

    add_bullet("Frontend Framework: ", "Next.js 16 (App Router) + React 19 + TypeScript. Server Components and Client Components mix for optimal performance.")
    add_bullet("Styling Engine: ", "Vanilla Tailwind CSS + globals.css custom Design System Tokens + CSS Variable Contrast Engine (Light/Dark mode support).")
    add_bullet("Database & ORM: ", "PostgreSQL database (Hosted on Neon Serverless Cloud) + Prisma ORM v5 (Type-safe query builder, migration engine, and schema definition).")
    add_bullet("Authentication & Security: ", "JWT (JSON Web Token) with HTTP-only Cookies and Bearer Header fallbacks + BcryptJS for password hashing (10 salt rounds).")
    add_bullet("CDN & Media Storage: ", "Cloudinary CDN with SHA1 signed API signature uploads via Node.js crypto module.")
    add_bullet("Data Visualization: ", "Recharts for dynamic line, area, bar, pie, and scatter charts.")

    add_h2("2.1 System Architecture Diagram")
    add_p("Niche di gayi diagram dikhati hai ki client browser se Request next.js server components, API route handlers, Prisma ORM, Neon PostgreSQL, aur Cloudinary CDN ke beech kaise flow karti hai:")

    if os.path.exists("docx_assets/arch_diagram.png"):
        doc.add_paragraph().paragraph_format.space_before = Pt(6)
        doc.add_picture("docx_assets/arch_diagram.png", width=Inches(6.2))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_c = p_cap.add_run("Figure 2.1: Kenzo OneERP Full Stack End-to-End System Architecture")
        r_c.font.size = Pt(9)
        r_c.font.italic = True
        r_c.font.color.rgb = RGBColor(100, 116, 139)

    add_callout(
        "Request Flow Step-by-Step:\n1. User browser client par request trigger karta hai (e.g. Leave apply or Profile Upload).\n2. Request Next.js App Router API Endpoint (/api/*) par land hoti hai.\n3. Middleware / getSession() function HTTP-only Cookie se JWT token extract aur verify karta hai.\n4. Route Handler Prisma Client (db.*) ke zariye Neon PostgreSQL DB se multi-tenant scoped SQL queries perform karta hai.\n5. Binary media uploads (Profile picture base64) Node.js 'crypto' SHA1 signed signature ke saath Cloudinary CDN API par POST hote hain.\n6. JSON response UI components ko pass hota hai jo real-time UI state re-render karte hain.",
        "DATA EXECUTION PIPELINE"
    )

    # --- CHAPTER 3: DATABASE SCHEMA & ENTITY RELATIONS ---
    add_h1("CHAPTER 3: Database Schema & Multi-Tenancy Data Model (Prisma + PostgreSQL)")

    add_p("Multi-tenancy ko implement karne ke liye Prisma schema mein har primary entity mein 'tenantId' field include kiya gaya hai. Primary entities mein Tenant reference constraints enforce hone se multi-company data leak bilkul impossible ho jata hai.")

    add_h2("3.1 Entity Relationship Diagram (ERD)")
    if os.path.exists("docx_assets/erd_diagram.png"):
        doc.add_picture("docx_assets/erd_diagram.png", width=Inches(6.2))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_c = p_cap.add_run("Figure 3.1: Database Entity Relationship Diagram (Prisma ORM Models)")
        r_c.font.size = Pt(9)
        r_c.font.italic = True
        r_c.font.color.rgb = RGBColor(100, 116, 139)

    add_h2("3.2 Prisma Schema Data Models Explanation")
    add_p("Here is the exact schema breakdown of all primary database models:")

    add_bullet("Tenant Model: ", "Multi-tenancy root container. Contains id (UUID), name, domain (unique), settings (JSON column for storing custom profiles, avatars, and branding).")
    add_bullet("User Model: ", "Authentication credential store. Fields: email (unique), passwordHash, name, role, tenantId. Linked to Tenant via Cascade Delete.")
    add_bullet("Employee Model: ", "Corporate Directory profile. Fields: firstName, lastName, department, position, status (ACTIVE/INACTIVE), salary (Float), hireDate. Linked 1-to-1 with User model via userId.")
    add_bullet("Attendance Model: ", "Daily check-in / check-out indexing. Fields: date (@db.Date), checkIn, checkOut, status (PRESENT/LATE/ABSENT/LEAVE). @@unique([employeeId, date]) constraint ensures duplicate attendance entries for the same date are impossible!")
    add_bullet("Leave Model: ", "Employee Leave Records. Fields: startDate, endDate, type (SICK/CASUAL/ANNUAL/UNPAID), status (PENDING/APPROVED/REJECTED), reason.")
    add_bullet("Project & Task Models: ", "Kanban Task Board. Project has budget (Float), status. Task has title, description, status (TODO/IN_PROGRESS/REVIEW/DONE), priority (LOW/MEDIUM/HIGH/URGENT/CRITICAL), assigneeId.")
    add_bullet("Invoice & Expense Models: ", "Financial Ledger. Invoice tracks amount in ₹, status (PAID/DRAFT). Expense tracks category, amount, status (PENDING/APPROVED/REJECTED), description.")
    add_bullet("AuditLog & Activity Models: ", "Security Auditing Ledger. Tracks action (LOGIN/CREATE/UPDATE/DELETE), resource, details, timestamp, and user reference.")

    add_code_block("""// Example Prisma Model Definition (prisma/schema.prisma)
model Attendance {
  id         String    @id @default(uuid())
  tenantId   String
  tenant     Tenant    @relation(fields: [tenantId], references: [id], onDelete: Cascade)
  employeeId String
  employee   Employee  @relation(fields: [employeeId], references: [id], onDelete: Cascade)
  date       DateTime  @db.Date
  checkIn    DateTime?
  checkOut   DateTime?
  status     String    // PRESENT, LATE, ABSENT, LEAVE
  createdAt  DateTime  @default(now())
  updatedAt  DateTime  @updatedAt

  @@unique([employeeId, date]) // Ensures no duplicate attendance records per employee per day
}""", "Prisma Schema")

    # --- CHAPTER 4: AUTHENTICATION & SECURITY ENGINE ---
    add_h1("CHAPTER 4: Authentication, Session Management & RBAC Security Engine")

    add_p("Security and Role-Based Access Control (RBAC) kisi bhi ERP system ka dil hota hai. KORE mein authentication Dual-Token verification system ke dwara handle ki jaati hai.")

    add_h2("4.1 Role-Based Access Control (RBAC) Pyramid")
    if os.path.exists("docx_assets/rbac_hierarchy.png"):
        doc.add_picture("docx_assets/rbac_hierarchy.png", width=Inches(6.0))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_c = p_cap.add_run("Figure 4.1: Role Hierarchy Pyramid in KORE")
        r_c.font.size = Pt(9)
        r_c.font.italic = True
        r_c.font.color.rgb = RGBColor(100, 116, 139)

    add_h2("4.2 Session Verification Engine (src/lib/auth.ts)")
    add_p("Next.js Server API Routes mein incoming requests ko verify karne ke liye `getSession(req)` function ka upyog kiya jata hai. Yeh function sabse pehle HTTP-only cookie 'token' check karta hai. Agar cookie absent ho, toh yeh Bearer Authorization Header inspect karta hai.")

    add_code_block("""// src/lib/auth.ts - Token Verification Logic
export async function getSession(req: NextRequest): Promise<SessionPayload | null> {
  // 1. Try Cookie first
  const cookieToken = req.cookies.get("token")?.value;
  if (cookieToken) {
    const payload = verifyToken(cookieToken);
    if (payload) return payload;
  }

  // 2. Try Authorization Header fallback
  const authHeader = req.headers.get("authorization");
  if (authHeader && authHeader.startsWith("Bearer ")) {
    const token = authHeader.substring(7);
    const payload = verifyToken(token);
    if (payload) return payload;
  }

  return null;
}""", "TypeScript - src/lib/auth.ts")

    add_callout(
        "User Profile Isolation Bug Solution:\nInitial implementation mein localStorage.getItem('kore_user_profile') key global thi, jiski wajah se Developer log in karne par Admin ka profile overwrite ho raha tha. Is bug ko humne local storage keys ko strictly user-scoped karke fix kiya: 'kore_user_profile_${user.id}'. Ab har account ka profile data aur Cloudinary avatar completely unique aur isolated rehta hai!",
        "CRITICAL BUG FIX INSIGHT"
    )

    # --- CHAPTER 5: CORE BUSINESS LOGICS & API ROUTES ---
    add_h1("CHAPTER 5: Core Business Logics & API Routes Walkthrough")

    add_p("Is section mein hum system ki sabse important business logics aur unke code implementations ko detail mein samjhenge.")

    add_h2("5.1 Cloudinary CDN Profile Upload Engine (`/api/profile`)")
    add_p("Jab user Edit Profile Modal se profile picture upload karta hai, toh Client browser Base64 Data URL generate karta hai. Server API `/api/profile` Node.js `crypto` module ka upyog karke SHA-1 signature compute karti hai aur directly Cloudinary CDN API par POST karti hai. Resulting Cloudinary CDN URL Database mein Tenant settings JSON object mein user ID ke against save hoti hai.")

    add_code_block("""// src/app/api/profile/route.ts - Cloudinary SHA1 Upload Logic
async function uploadToCloudinary(base64Data: string) {
  const cloudName = process.env.CLOUD_NAME || "dg7tgftmf";
  const apiKey = process.env.CLOUDINARY_API_KEY || "375739656182662";
  const apiSecret = process.env.CLOUDINARY_API_SECRET || "DIzbPW9-zWGsPERPJnHtgElutG4";

  const timestamp = Math.floor(Date.now() / 1000).toString();
  const folder = "kore_profiles";
  const transformation = "c_fill,g_face,w_300,h_300,q_auto,f_auto";

  // Create SHA1 signature for Cloudinary API security
  const strToSign = `folder=${folder}&transformation=${transformation}&timestamp=${timestamp}${apiSecret}`;
  const signature = crypto.createHash("sha1").update(strToSign).digest("hex");

  const formData = new FormData();
  formData.append("file", base64Data);
  formData.append("api_key", apiKey);
  formData.append("timestamp", timestamp);
  formData.append("folder", folder);
  formData.append("transformation", transformation);
  formData.append("signature", signature);

  const res = await fetch(`https://api.cloudinary.com/v1_1/${cloudName}/image/upload`, {
    method: "POST",
    body: formData,
  });

  const data = await res.json();
  if (data.secure_url) return data.secure_url;
  throw new Error("Cloudinary upload failed");
}""", "TypeScript - Cloudinary Engine")

    add_h2("5.2 Real-Time Leave Application & Isolation Engine (`/api/leaves`)")
    if os.path.exists("docx_assets/dfd_level1.png"):
        doc.add_picture("docx_assets/dfd_level1.png", width=Inches(6.0))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_c = p_cap.add_run("Figure 5.1: Real-time Leave Flow & Data Isolation Pipeline")
        r_c.font.size = Pt(9)
        r_c.font.italic = True
        r_c.font.color.rgb = RGBColor(100, 116, 139)

    add_p("Leave Management Module do strict rules follow karta hai:")
    add_bullet("1. Real-Time Approval Engine: ", "Employee dwara submit ki gayi leave request instantly PENDING status mein Database mein write hoti hai. Admin, CEO, aur HR Dashboards real-time background polling / state re-fetch ke dwara is request ko view aur instantly APPROVED ya REJECTED mark kar sakte hain.")
    add_bullet("2. Role-Scoped Leave Isolation & Self-Deletion: ", "Employee sirf aur sirf apni khud ki leaves dekh sakta hai aur unhe delete kar sakta hai (`DELETE /api/leaves?id=...`). Regular employees doosre colleagues ki leaves nahi dekh sakte. Tabhi Admin, CEO, aur HR roles pooray tenant ki sabhi leave applications view aur delete kar sakte hain.")

    if os.path.exists("docx_assets/leave_workflow.png"):
        doc.add_picture("docx_assets/leave_workflow.png", width=Inches(6.0))

    add_h2("5.3 Task Kanban Board & Top-Bar Notification Engine (`/api/projects/tasks`)")
    add_p("Projects Kanban board par Jab Admin, CEO, ya HR kisi Employee ko Naya task assign karta hai, toh frontend API ko POST request bhejtaye hain. Success response aate hi Top Bar par Real-Time Toast Banner ('Task assigned successfully') trigger hota hai.")

    # --- CHAPTER 6: THEME SYSTEM & UI CONTRAST ENGINE ---
    add_h1("CHAPTER 6: Theme System, UI Contrast Engine & Styling Architecture")

    add_p("KORE system Light Mode aur Dark Mode dono ko seamlessly support karta hai. Initial stages mein theme switching ke dauran text visibility bugs (black-on-black text ya white-on-white text) ko eliminate karne ke liye humne ek robust CSS Contrast Engine design kiya.")

    if os.path.exists("docx_assets/theme_engine.png"):
        doc.add_picture("docx_assets/theme_engine.png", width=Inches(5.8))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_c = p_cap.add_run("Figure 6.1: Theme Engine & Contrast Rules Matrix")
        r_c.font.size = Pt(9)
        r_c.font.italic = True
        r_c.font.color.rgb = RGBColor(100, 116, 139)

    add_h2("6.1 Light & Dark Contrast Rules (src/app/globals.css)")
    add_code_block("""/* Light Mode Text Contrast Engine */
[data-theme="light"] body,
[data-theme="light"] table,
[data-theme="light"] .glass-panel,
[data-theme="light"] p,
[data-theme="light"] label {
  color: #0f172a; /* Crisp dark slate text */
}

/* Dark Mode Text Contrast Engine */
[data-theme="dark"] body,
[data-theme="dark"] table,
[data-theme="dark"] .glass-panel,
[data-theme="dark"] p,
[data-theme="dark"] label {
  color: #f8fafc; /* Pure bright white text */
}

/* Fixed Profile Modal Dark Theme Isolation Rule */
.profile-modal-dark,
[data-theme="light"] .profile-modal-dark {
  background-color: #0f172a !important;
  color: #ffffff !important;
}

.profile-modal-dark input {
  background-color: #1e293b !important;
  color: #ffffff !important;
  -webkit-text-fill-color: #ffffff !important;
}""", "CSS - src/app/globals.css")

    add_callout(
        "Profile Modal Theme Isolation (.profile-modal-dark):\nUser ki request ke according Edit Profile Modal ko universal dark glassmorphic styling mein freeze kiya gaya hai. Matlab chahe poori site Light Theme mein ho ya Dark Theme mein, Edit Profile Modal humesha same dark sleek design (#0f172a) mein high-visibility white typed text (#ffffff) ke saath render hoga!",
        "THEME IMMUNITY DESIGN PATTERN"
    )

    # --- CHAPTER 7: FRONTEND PAGE & COMPONENT ARCHITECTURE ---
    add_h1("CHAPTER 7: Frontend Page & Component Architecture Walkthrough")

    add_p("Frontend Codebase modular Architecture follow karta hai. Essential pages aur components ka structure neeche explain kiya gaya hai:")

    add_h2("7.1 Main Dashboard Router (`src/app/(dashboard)/dashboard/page.tsx`)")
    add_p("Dashboard main page current logged-in user ke Role ke according dynamic view render karta hai:")
    add_bullet("Admin View: ", "User management, password force-reset modal, spawn new user instance, system audit logs.")
    add_bullet("CEO Intelligence View: ", "Financial capital inflow, burn rate trends, deals closed, Recharts revenue area chart.")
    add_bullet("HR Insights View: ", "Attendance index donut chart, leave approval queue, out-of-office calendar.")
    add_bullet("Developer / Employee Workspace View: ", "Personal ongoing tasks volume, active projects progress, budget lines.")

    add_h2("7.2 Additional Dashboard Module Pages")
    add_bullet("analytics/page.tsx: ", "Deep financial reports, headcount pie charts, target revenue bar charts, expense distributions.")
    add_bullet("companies/page.tsx: ", "Corporate tenant overview, clean borderless company logo display (`/logo.png`), active subscription tier.")
    add_bullet("crm/page.tsx: ", "Interactive lead management pipeline, conversion funnel visualization, total deal value metrics in ₹.")
    add_bullet("departments/page.tsx: ", "Department cards (Engineering, HR, Finance, Sales), headcount metrics, average salary in ₹, team avatar click-to-view profiles.")
    add_bullet("employees/page.tsx: ", "Comprehensive employee directory table, role filter badges, search query filter, salary and position updates.")
    add_bullet("finance/page.tsx: ", "Invoices ledger, spending breakdown, expense approval workflow.")
    add_bullet("hrms/page.tsx: ", "HR operations portal, attendance check-in button, leave self-deletion history table.")
    add_bullet("projects/page.tsx: ", "Kanban drag-and-drop board with TODO, IN_PROGRESS, REVIEW, DONE columns.")
    add_bullet("copilot/page.tsx: ", "AI Assistant chat interface with quick action presets.")

    # --- CHAPTER 8: PRACTICAL SETUP & SENIOR DEVELOPER GUIDELINES ---
    add_h1("CHAPTER 8: Practical Setup, Deployment & Senior Developer Rules")

    add_p("Is final chapter mein local setup steps, production build validation, aur senior software engineering best practices detailed tarike se diye gaye hain.")

    add_h2("8.1 Local Environment & Deployment Workflow")
    add_bullet("1. Clone Repository & Install Dependencies: ", "git clone <repo-url> && npm install")
    add_bullet("2. Environment Variables (.env): ", "Configure DATABASE_URL (Neon PostgreSQL), JWT_SECRET, CLOUD_NAME, CLOUDINARY_API_KEY, CLOUDINARY_API_SECRET.")
    add_bullet("3. Database Synchronization: ", "npx prisma generate && npx prisma db push --accept-data-loss")
    add_bullet("4. Start Development Server: ", "npm run dev (Runs on http://localhost:3000)")
    add_bullet("5. Production Build Verification: ", "npm run build (Ensures 0 TypeScript errors & 0 build warnings)")
    add_bullet("6. Git Version Control Push: ", "git add . && git commit -m 'feat: ...' && git push origin main")

    add_h2("8.2 Top 10 Senior Developer Commandments")
    add_p("Aapko ek expert ERP System Specialist banne ke liye always in 10 commandments ko obey karna chahiye:")
    add_bullet("1. Never Guess Schema Paths: ", "Humesha prisma/schema.prisma aur API route definitions inspect karein code likhne se pehle.")
    add_bullet("2. User-Scoped Local State: ", "Local storage keys ko humesha user ID se prefix karein (e.g. kore_user_profile_${user.id}) data contamination rokne ke liye.")
    add_bullet("3. Strict Multi-Tenancy Scoping: ", "Database queries mein tenantId WHERE clause kabhi na bhulein.")
    add_bullet("4. Empirical Build Verification: ", "Code edit karne ke baad bina 'npm run build' run kiye kabhi success declare na karein.")
    add_bullet("5. Silent Log Inspection: ", "Errors hone par raw terminal stack traces padhein, blind guess na karein.")
    add_bullet("6. Theme Immunity Patterns: ", "Fixed-styled components (like profile modals) ke liye explicitly scoped CSS classes (.profile-modal-dark) use karein.")
    add_bullet("7. Currency Standardisation: ", "Monetary metrics mein standard currency symbols (₹) maintain karein.")
    add_bullet("8. Secure Media Storage: ", "Base64 images directly DB mein store karne ke bajaye Cloudinary CDN signed uploads use karein.")
    add_bullet("9. Role Privilege Validation: ", "Backend route handlers mein session.role verify karein client-side checks ke alaawa.")
    add_bullet("10. Clean Code & Documentation: ", "Self-documenting code likhein aur maintainable project architecture follow karein.")

    # Save document
    filename = "Kenzo_OneERP_Masterclass_Guide.docx"
    doc.save(filename)
    print(f"Masterclass document successfully created: {filename}")

if __name__ == "__main__":
    create_document()
